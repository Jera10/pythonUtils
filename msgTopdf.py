# Installation:
# pip install extract-msg python-docx docx2pdf

import os
import sys
import extract_msg
import re
from docx import Document
from docx2pdf import convert

def sanitize_filename(filename):
    return re.sub(r'[<>:"/\\\\|?*]', '_', filename)

def convert_msg_to_docx(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith(".msg"):
            msg_path = os.path.join(input_folder, filename)
            msg = extract_msg.Message(msg_path)
            subject = msg.subject or "No_Subject"
            body = msg.body or "No_Content"

            sanitized_subject = sanitize_filename(subject)
            docx_path = os.path.join(output_folder, sanitized_subject + ".docx")

            doc = Document()
            doc.add_heading(subject, 0)
            doc.add_paragraph(body)
            doc.save(docx_path)

            print(f"Converted: {msg_path} → {docx_path}")

def convert_docx_to_pdf(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    for filename in os.listdir(input_folder):
        if filename.endswith(".docx"):
            docx_path = os.path.join(input_folder, filename)
            pdf_path = os.path.join(output_folder, os.path.splitext(filename)[0] + ".pdf")
            convert(docx_path, pdf_path)
            print(f"Converted: {docx_path} → {pdf_path}")

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python script_name.py <input_folder> <output_folder>")
        sys.exit(1)

    input_folder = sys.argv[1]
    output_folder = sys.argv[2]

    convert_msg_to_docx(input_folder, output_folder)
    convert_docx_to_pdf(output_folder, output_folder)
